# -*- coding: utf-8 -*-
import os
import logging
from utils import merge_audio_chunks_direct, build_docx_and_pdf, build_transcript_from_cache
from datetime import datetime
import threading
import queue

# Setup logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

MEETINGS_DIR = os.getenv("MEETINGS_DIR", "meetings")

# Global job queue
job_queue = queue.Queue()

class JobWorker(threading.Thread):
    """Background worker thread"""
    def __init__(self):
        super().__init__(daemon=True)
        self.running = True

    def run(self):
        logger.info("🔄 Job Worker started")
        while self.running:
            try:
                # Lấy job từ queue
                job_type, args, kwargs = job_queue.get(timeout=1)
                logger.info(f"⚙️ Processing job: {job_type} with args={args}")

                # Xử lý từng loại job theo thứ tự
                if job_type == "stt":
                    result = enqueue_stt_job(*args, **kwargs)
                elif job_type == "merge_transcript":
                    # Đảm bảo tất cả các job STT liên quan đã hoàn thành
                    meeting_id = args[0]
                    logger.info(f"Waiting for all STT jobs to complete for meeting_id={meeting_id}")
                    wait_for_stt_jobs(meeting_id)
                    result = enqueue_merge_transcript_job(*args, **kwargs)
                elif job_type == "merge_audio":
                    result = enqueue_merge_job(*args, **kwargs)

                logger.info(f"✅ Job completed: {result}")
                job_queue.task_done()

            except queue.Empty:
                continue
            except Exception as e:
                logger.error(f"❌ Job failed: {str(e)}", exc_info=True)
                job_queue.task_done()

    def stop(self):
        self.running = False

# Start worker threads
NUM_WORKERS = 2  # Số lượng worker (tùy chỉnh theo số CPU)
workers = [JobWorker() for _ in range(NUM_WORKERS)]
for worker in workers:
    worker.start()

def enqueue_job(job_type, *args, **kwargs):
    """Enqueue job to be processed by worker threads"""
    job_queue.put((job_type, args, kwargs))
    logger.info(f"📥 Job enqueued: {job_type}")

# Correct the import to use the existing function in utils.py
from utils import get_whisper_model

# Initialize Whisper model globally
# Load the Whisper model once at the start
WHISPER_MODEL = get_whisper_model()

def enqueue_stt_job(meeting_id, user_id, full_name, role, ts, filepath):
    """
    Transcribe audio file using Whisper and store result in Redis cache
    """
    try:
        logger.info(f"Starting STT job for meeting_id={meeting_id}, user_id={user_id}, file={filepath}")

        from utils import transcribe_with_whisper, append_transcript_cache

        # Transcribe audio using the cached Whisper model
        logger.info(f"Transcribing {filepath}...")
        text = transcribe_with_whisper(filepath)
        logger.info(f"Transcription complete. Text length: {len(text)}")

        # Create entry
        entry = {
            "ts": ts,
            "user_id": user_id,
            "full_name": full_name,
            "role": role,
            "text": text,
            "source_file": filepath
        }

        # Append to cache
        logger.info(f"Appending to cache for meeting_id={meeting_id}")
        append_transcript_cache(meeting_id, entry)

        result = {"meeting_id": meeting_id, "user_id": user_id, "text_len": len(text)}
        logger.info(f"STT job completed successfully: {result}")
        return result

    except Exception as e:
        logger.error(f"STT job failed: {str(e)}", exc_info=True)
        raise RuntimeError(f"STT job failed for meeting_id={meeting_id}, user_id={user_id}: {str(e)}")


def enqueue_merge_transcript_job(meeting_id):
    """
    Merge all transcripts from Redis cache and create DOCX file.
    Wait for all STT jobs to complete before proceeding.
    Delete old transcript files before creating new ones.
    """
    try:
        logger.info(f"Starting merge transcript job for meeting_id={meeting_id}")

        # Wait for all STT jobs to complete
        from utils import wait_for_stt_jobs, delete_old_transcripts
        wait_for_stt_jobs(meeting_id)

        meeting_dir = os.path.join(MEETINGS_DIR, meeting_id)
        final_dir = os.path.join(meeting_dir, "final")
        timestamp = datetime.utcnow().strftime("%d-%m-%Y_%H-%M-%S")

        os.makedirs(final_dir, exist_ok=True)

        # Delete old transcript files
        logger.info(f"Deleting old transcript files in {final_dir}")
        delete_old_transcripts(final_dir)

        # Get transcripts from cache
        logger.info(f"Fetching transcripts from cache for meeting_id={meeting_id}")
        entries = build_transcript_from_cache(meeting_id)
        logger.info(f"Found {len(entries)} transcripts in cache")

        if not entries:
            raise RuntimeError("No transcripts found in cache")

        # Create DOCX file
        docx_path = os.path.join(final_dir, f"transcript_{meeting_id}_{timestamp}.docx")
        logger.info(f"Creating DOCX file: {docx_path}")

        from docx import Document
        doc = Document()
        doc.add_heading(f"Bien ban cuoc hop: {meeting_id}", level=1)
        doc.add_paragraph(f"Created: {datetime.utcnow().strftime('%d/%m/%Y %H:%M:%S UTC')}")
        doc.add_paragraph("")

        for e in entries:
            ts_str = e.get("ts", "")
            full_name = e.get("full_name", "Unknown")
            role = e.get("role", "")
            text = e.get("text", "")
            line = f"({ts_str}) {full_name} - {role}: {text}"
            doc.add_paragraph(line)

        doc.save(docx_path)
        logger.info(f"DOCX file saved successfully: {docx_path}")

        result = {
            "status": "transcript_created",
            "meeting_id": meeting_id,
            "output": docx_path,
            "total_transcripts": len(entries)
        }
        logger.info(f"Merge transcript job completed: {result}")
        return result

    except Exception as e:
        logger.error(f"Merge transcript job failed: {str(e)}", exc_info=True)
        raise RuntimeError(f"Merge transcript failed for meeting_id={meeting_id}: {str(e)}")


def enqueue_merge_job(meeting_id):
    meeting_dir = os.path.join(MEETINGS_DIR, meeting_id)
    chunks_dir = os.path.join(meeting_dir, "chunks")
    final_dir = os.path.join(meeting_dir, "final")
    log_path = os.path.join(meeting_dir, "merge.log")
    timestamp = datetime.utcnow().strftime("%d-%m-%Y_%H-%M-%S")
    merged_ogg_path = None
    
    try:
        os.makedirs(final_dir, exist_ok=True)
        
        with open(log_path, "a", encoding="utf-8") as log:
            log.write(f"\n=== Merge started at {datetime.utcnow()} ===\n")
            log.write(f"Meeting ID: {meeting_id}\n")
            log.write(f"Chunks dir: {chunks_dir}\n")
            log.write(f"Final dir: {final_dir}\n")
            log.flush()
            
            try:
                if not os.path.exists(chunks_dir):
                    raise RuntimeError(f"Chunks directory does not exist: {chunks_dir}")
                
                if not os.listdir(chunks_dir):
                    raise RuntimeError("No audio chunks to merge")
                
                merged_ogg_path = os.path.join(final_dir, f"merged_{timestamp}.ogg")
                log.write(f"Output OGG file: {merged_ogg_path}\n")
                log.flush()
                
                try:
                    old_ogg_files = [f for f in os.listdir(final_dir) if f.endswith(".ogg")]
                    deleted_count = 0
                    for old_ogg_file in old_ogg_files:
                        old_ogg_path = os.path.join(final_dir, old_ogg_file)
                        try:
                            os.remove(old_ogg_path)
                            deleted_count += 1
                            log.write(f"Deleted old OGG file: {old_ogg_file}\n")
                        except Exception as e:
                            log.write(f"Failed to delete {old_ogg_file}: {e}\n")
                    log.write(f"Total deleted old OGG files: {deleted_count}/{len(old_ogg_files)}\n")
                    log.flush()
                except Exception as e:
                    log.write(f"Error deleting old OGG files: {e}\n")
                    log.flush()

                log.write("Starting audio merge with direct format conversion...\n")
                log.flush()
                merge_audio_chunks_direct(chunks_dir, merged_ogg_path, log_file=log_path)
                log.write(f"Merge and OGG conversion completed successfully!\n")
                log.write(f"Merged OGG file: {merged_ogg_path}\n")
                log.flush()

            except Exception as e:
                log.write(f"Merge failed: {e}\n")
                log.write(f"Error type: {type(e).__name__}\n")
                import traceback
                log.write(f"Traceback: {traceback.format_exc()}\n")
                log.flush()
                raise

            finally:
                log.write(f"=== Merge ended at {datetime.utcnow()} ===\n")
                log.flush()

    except Exception as e:
        try:
            with open(log_path, "a", encoding="utf-8") as log:
                log.write(f"\n=== CRITICAL ERROR at {datetime.utcnow()} ===\n")
                log.write(f"Error: {e}\n")
                log.write(f"Error type: {type(e).__name__}\n")
                import traceback
                log.write(f"Traceback: {traceback.format_exc()}\n")
                log.flush()
        except:
            pass
        raise

    return {
        "status": "merged",
        "output": merged_ogg_path if merged_ogg_path and os.path.exists(merged_ogg_path) else None
    }
