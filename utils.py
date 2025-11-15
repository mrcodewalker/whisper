# -*- coding: utf-8 -*-
import os, json, hashlib, shutil, wave
from redis import Redis
from datetime import datetime
from docx import Document
from pydub import AudioSegment

REDIS_URL = os.getenv("REDIS_URL", "redis://localhost:6379/0")
r = Redis.from_url(REDIS_URL)
MEETINGS_DIR = os.getenv("MEETINGS_DIR", "meetings")

# Global model cache
_whisper_model = None

def get_whisper_model(model_name="medium"):
    """
    Get cached Whisper model to avoid reloading every time
    """
    global _whisper_model
    if _whisper_model is None:
        try:
            import whisper
            print(f"Loading Whisper model '{model_name}'...")
            _whisper_model = whisper.load_model(model_name)
            print(f"Model loaded successfully!")
        except Exception as e:
            raise RuntimeError("Failed to load whisper model: " + str(e))
    return _whisper_model

def merge_audio_chunks_direct(chunks_dir, out_path, log_file=None):
    """
    Merge all audio files (.wav, .ogg, .m4a, etc.) directly to OGG format using pydub.
    This avoids ffmpeg concat issues with opus codec.
    
    Args:
        chunks_dir: Directory containing audio chunks
        out_path: Output file path (should end with .ogg)
        log_file: Optional log file path to write detailed logs
    """
    def log_msg(msg):
        print(msg)
        if log_file:
            try:
                with open(log_file, "a", encoding="utf-8") as f:
                    f.write(msg + "\n")
                    f.flush()
            except:
                pass
    
    if not os.path.exists(chunks_dir):
        raise RuntimeError("Audio chunks directory does not exist")
    
    # Get all audio files
    audio_files = []
    for f in os.listdir(chunks_dir):
        if f.lower().endswith(('.wav', '.ogg', '.mp3', '.m4a', '.flac', '.opus')):
            audio_files.append(f)
    
    if not audio_files:
        raise RuntimeError("No audio chunks to merge")
    
    # Sort by timestamp
    def extract_timestamp(filename):
        try:
            date_part = filename.split("__")[0]
            dt = datetime.strptime(date_part, "%d-%m-%Y_%H-%M-%S")
            return dt
        except:
            return datetime.min
    
    audio_files.sort(key=extract_timestamp)
    log_msg(f"Found {len(audio_files)} audio files to merge")
    
    # Merge all audio files
    merged_audio = None
    successful_merges = 0
    
    for i, fname in enumerate(audio_files):
        fpath = os.path.join(chunks_dir, fname)
        try:
            log_msg(f"Processing {i+1}/{len(audio_files)}: {fname}")
            
            # Load audio with pydub (supports multiple formats)
            audio = AudioSegment.from_file(fpath)
            
            if merged_audio is None:
                merged_audio = audio
                log_msg(f"  -> Initialized with {audio.duration_seconds:.2f}s ({audio.channels}ch, {audio.frame_rate}Hz)")
            else:
                # Make sure same format before concatenating
                if audio.frame_rate != merged_audio.frame_rate:
                    audio = audio.set_frame_rate(merged_audio.frame_rate)
                if audio.channels != merged_audio.channels:
                    audio = audio.set_channels(merged_audio.channels)
                
                merged_audio += audio
                log_msg(f"  -> Added {audio.duration_seconds:.2f}s, total now: {merged_audio.duration_seconds:.2f}s")
            
            successful_merges += 1
            
        except Exception as e:
            log_msg(f"  -> WARNING: Failed to process {fname}: {str(e)}")
            continue
    
    if merged_audio is None:
        raise RuntimeError("Failed to merge any audio files")
    
    log_msg(f"\nSuccessfully merged {successful_merges}/{len(audio_files)} files")
    log_msg(f"Total duration: {merged_audio.duration_seconds:.2f} seconds")
    
    # Export to OGG format
    try:
        log_msg(f"\nExporting to OGG format: {out_path}")
        # Create output directory if needed
        os.makedirs(os.path.dirname(out_path) if os.path.dirname(out_path) else ".", exist_ok=True)
        merged_audio.export(out_path, format="ogg", bitrate="128k", codec="libvorbis", parameters=["-q:a", "7"])
        log_msg(f"OGG export completed successfully!")
        log_msg(f"Output file size: {os.path.getsize(out_path) / (1024*1024):.2f} MB")
        return out_path
        
    except Exception as e:
        log_msg(f"ERROR during OGG export: {str(e)}")
        if os.path.exists(out_path):
            os.remove(out_path)
        raise RuntimeError(f"Failed to export OGG: {str(e)}")


def merge_audio_chunks(chunks_dir, out_path):
    if not os.path.exists(chunks_dir):
        raise RuntimeError("Audio chunks directory does not exist")
    
    files = [f for f in os.listdir(chunks_dir) if f.endswith(".wav")]
    if not files:
        raise RuntimeError("No audio chunks to merge")
    
    def extract_timestamp(filename):
        try:
            date_part = filename.split("__")[0]  # Lấy phần timestamp
            dt = datetime.strptime(date_part, "%d-%m-%Y_%H-%M-%S")
            return dt
        except:
            return datetime.min  # File không đúng format sẽ đẩy lên đầu
    
    files.sort(key=extract_timestamp)
    
    n_channels = None
    sampwidth = None
    framerate = None
    valid_files = []
    
    print(f"Checking {len(files)} files...")
    
    for fname in files:
        fpath = os.path.join(chunks_dir, fname)
        try:
            with wave.open(fpath, 'rb') as wav_file:
                if n_channels is None:
                    params = wav_file.getparams()
                    n_channels = params.nchannels
                    sampwidth = params.sampwidth
                    framerate = params.framerate
                    print(f"Using audio params from {fname}: {n_channels}ch, {sampwidth}bytes, {framerate}Hz")
                valid_files.append(fname)
        except Exception as e:
            print(f"WARNING: Skipping corrupted file {fname}: {str(e)}")
            continue
    
    if not valid_files:
        raise RuntimeError("No valid WAV files found to merge")
    
    print(f"Found {len(valid_files)} valid files out of {len(files)}")
    
    try:
        with wave.open(out_path, 'wb') as output:
            output.setnchannels(n_channels)
            output.setsampwidth(sampwidth)
            output.setframerate(framerate)
            
            merged_count = 0
            for i, fname in enumerate(valid_files):
                fpath = os.path.join(chunks_dir, fname)
                print(f"Merging {i+1}/{len(valid_files)}: {fname}")
                
                try:
                    with wave.open(fpath, 'rb') as wav_file:
                        if (wav_file.getnchannels() != n_channels or
                            wav_file.getsampwidth() != sampwidth or
                            wav_file.getframerate() != framerate):
                            print(f"WARNING: {fname} has different audio parameters, skipping...")
                            continue
                        
                        frames = wav_file.readframes(wav_file.getnframes())
                        output.writeframes(frames)
                        merged_count += 1
                        
                except Exception as e:
                    print(f"ERROR reading {fname}: {str(e)}, skipping...")
                    continue
        
        print(f"Successfully merged {merged_count}/{len(files)} files to {out_path}")
        return out_path
        
    except Exception as e:
        if os.path.exists(out_path):
            os.remove(out_path)
        raise RuntimeError(f"Error creating merged file: {str(e)}")


def transcribe_with_whisper(filepath):
    """
    Transcribe audio file using Whisper model with caching
    """
    try:
        import whisper
    except Exception as e:
        raise RuntimeError("whisper package not installed: " + str(e))
    
    model = get_whisper_model("base")
    result = model.transcribe(filepath)
    return result.get("text", "").strip()

def append_transcript_cache(meeting_id, entry):
    cache_key = f"meeting:{meeting_id}:transcripts"
    last = r.lindex(cache_key, -1)
    if last:
        last_e = json.loads(last)
        try:
            fmt_ddmmyyyy = "%d-%m-%Y_%H-%M-%S"
            tlast = datetime.strptime(last_e["ts"], fmt_ddmmyyyy)
            tcur = datetime.strptime(entry["ts"], fmt_ddmmyyyy)
            gap = (tcur - tlast).total_seconds()
        except Exception:
            gap = 9999
        if last_e.get("user_id") == entry.get("user_id") and gap <= 30:
            last_e["text"] = last_e.get("text", "") + " " + entry.get("text", "")
            r.rpop(cache_key)
            r.rpush(cache_key, json.dumps(last_e))
            return
    r.rpush(cache_key, json.dumps(entry))

def build_docx_and_pdf(meeting_id, entries, output_dir):
    doc = Document()
    doc.add_heading(f"Bien ban cuoc hop: {meeting_id}", level=1)
    doc.add_paragraph(f"Created: {datetime.utcnow().strftime('%d/%m/%Y %H:%M:%S UTC')}")
    doc.add_paragraph("")
    for e in entries:
        ts_str = e.get("ts", "")
        line = f"({ts_str}) {e.get('full_name','Unknown')} - {e.get('role','')}: {e.get('text','')}"
        doc.add_paragraph(line)
    docx_path = os.path.join(output_dir, f"{meeting_id}.docx")
    doc.save(docx_path)
    pdf_path = os.path.join(output_dir, f"{meeting_id}.pdf")
    try_convert_docx_to_pdf_libreoffice(docx_path, pdf_path)
    return docx_path, pdf_path

def try_convert_docx_to_pdf_libreoffice(docx_path, pdf_path):
    try:
        outdir = os.path.dirname(pdf_path)
        import subprocess
        cmd = ["soffice", "--headless", "--convert-to", "pdf", "--outdir", outdir, docx_path]
        subprocess.check_call(cmd, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
        produced = os.path.join(outdir, os.path.splitext(os.path.basename(docx_path))[0] + ".pdf")
        if os.path.exists(produced):
            if produced != pdf_path:
                os.replace(produced, pdf_path)
            return True
    except Exception:
        return False
    return False

def build_transcript_from_cache(meeting_id):
    cache_key = f"meeting:{meeting_id}:transcripts"
    entries = r.lrange(cache_key, 0, -1)
    return [json.loads(e) for e in entries]

def clear_transcript_cache(meeting_id):
    """
    Clear transcript cache for a specific meeting ID.
    """
    try:
        cache_key = f"meeting:{meeting_id}:transcripts"
        r.delete(cache_key)  # Assuming `r` is the Redis connection
        print(f"✅ Cleared transcript cache for meeting_id={meeting_id}")
    except Exception as e:
        print(f"❌ Failed to clear transcript cache for meeting_id={meeting_id}: {str(e)}")